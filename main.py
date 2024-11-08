import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import io
from PIL import Image

def extract_text_with_pypdf2(pdf_path):
    text_content = []
    with open(pdf_path, "rb") as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        for page_num, page in enumerate(pdf_reader.pages):
            text = page.extract_text()
            text_content.append(f"Page {page_num + 1}:\n{text}\n" if text else f"Page {page_num + 1}: No text found.\n")
    return text_content

def extract_images_with_pymupdf(pdf_path):
    images_content = []
    pdf_document = fitz.open(pdf_path)

    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        page_images = []
        
        for img_index, img in enumerate(page.get_images(full=True)):
            xref = img[0]  # The image reference number in the PDF

            # Extract image bytes from the PDF
            base_image = pdf_document.extract_image(xref)
            img_bytes = base_image["image"]

            # Use Pillow to convert the image to a PNG format in a BytesIO stream
            img_stream = io.BytesIO(img_bytes)
            pil_image = Image.open(img_stream).convert("RGB")  # Ensure it's in RGB mode
            png_stream = io.BytesIO()
            pil_image.save(png_stream, format="PNG")  # Convert to PNG
            png_stream.seek(0)

            page_images.append(png_stream)  # Store the PNG stream
        images_content.append(page_images)

    pdf_document.close()
    return images_content

def pdf_to_word(pdf_path, word_path):
    # Create a new Word document
    doc = Document()
    doc.add_heading("PDF to Word Conversion", level=1)

    # Extract text using PyPDF2
    text_content = extract_text_with_pypdf2(pdf_path)
    
    # Extract images using PyMuPDF
    images_content = extract_images_with_pymupdf(pdf_path)

    # Add extracted text and images to the Word document
    for page_num, page_text in enumerate(text_content):
        doc.add_heading(f"Page {page_num + 1}", level=2)
        doc.add_paragraph(page_text)

        # Add images if present
        if page_num < len(images_content) and images_content[page_num]:
            for img_index, img_stream in enumerate(images_content[page_num]):
                img_paragraph = doc.add_paragraph()
                img_paragraph.add_run(f"Image {img_index + 1} on Page {page_num + 1}").bold = True
                doc.add_picture(img_stream, width=Inches(5))  # Adjust size as needed
                doc.add_paragraph("")

        # Add a page break after each page
        doc.add_page_break()

    # Save the Word document
    doc.save(word_path)
    print(f"Conversion complete! The Word document is saved at: {word_path}")

# Run the function with the paths to your PDF and the desired output Word file
pdf_to_word("input_pdf_file.pdf", "output_word_file.docx")
